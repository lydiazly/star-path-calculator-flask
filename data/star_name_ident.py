import pandas as pd
from docx import Document

doc = Document('ident6.docx')
# doc = Document('ident4.docx')
full_text = []
for para in doc.paragraphs:
    full_text.append(para.text)
    
df = pd.DataFrame(columns={"Proper Name": str, "HIP": int})
# df = pd.DataFrame(columns={"Bayer Designation": str, "HIP": int})

for i in range(len(full_text)):
    a, b = full_text[i].split('|')
    df.loc[i] = [a.strip(), int(b.strip())]

df.to_pickle("ident_proper.pkl")
# df.to_pickle("ident_bayer.pkl")